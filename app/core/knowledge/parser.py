"""文档解析器"""
import os
from typing import Optional
from pathlib import Path

# 尝试导入 PyMuPDF
try:
    import fitz
except ImportError:
    pass


class DocumentParser:
    """文档解析器"""
    
    def parse(self, file_path: str) -> str:
        """
        解析文档文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析后的文本内容
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension == ".md":
            return self._parse_markdown(file_path)
        elif file_extension == ".txt":
            return self._parse_text(file_path)
        elif file_extension == ".pdf":
            return self._parse_pdf(file_path)
        elif file_extension == ".docx":
            return self._parse_word(file_path)
        elif file_extension == ".json":
            return self._parse_json(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
    
    def _parse_markdown(self, file_path: Path) -> str:
        """解析 Markdown 文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            raise Exception(f"解析 Markdown 文件时出错: {str(e)}")
    
    def _parse_text(self, file_path: Path) -> str:
        """解析文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            raise Exception(f"解析文本文件时出错: {str(e)}")
    
    def _parse_pdf(self, file_path: Path) -> str:
        """解析 PDF 文件"""
        try:
            import fitz  # PyMuPDF
            text = ""
            
            with fitz.open(file_path) as pdf:
                for page_num, page in enumerate(pdf):
                    page_text = ""
                    
                    try:
                        # 提取页面的视觉块
                        blocks = self._extract_visual_blocks(page)
                        
                        # 处理视觉块
                        processed_blocks = self._process_visual_blocks(blocks, page_num + 1, file_path)
                        
                        # 按阅读顺序重组文本
                        ordered_content = self._reorder_content(processed_blocks)
                        
                        # 合并处理结果
                        page_text = "\n".join(ordered_content)
                        
                    except Exception as page_error:
                        page_text = f"[第{page_num + 1}页解析失败: {str(page_error)}]\n"
                    
                    text += page_text + "\n"
            
            return text
        except ImportError:
            raise Exception("解析 PDF 文件需要 PyMuPDF 库，请安装: pip install PyMuPDF")
        except Exception as e:
            raise Exception(f"解析 PDF 文件时出错: {str(e)}")
    
    def _extract_visual_blocks(self, page) -> list:
        """提取页面的视觉块（文本区、表格区、图片区）"""
        try:
            # 使用 PyMuPDF 提取页面的所有块
            blocks = page.get_text("blocks")
            
            # 分类块为文本、表格、图片
            visual_blocks = []
            for block in blocks:
                x0, y0, x1, y1, text, block_type, block_number = block
                
                block_info = {
                    "type": "text",
                    "bbox": (x0, y0, x1, y1),
                    "content": text.strip(),
                    "page_num": page.number + 1
                }
                
                # 简单判断：如果块中包含大量表格特征，可能是表格
                if self._is_table_candidate(text):
                    block_info["type"] = "table_candidate"
                
                visual_blocks.append(block_info)
            
            # 检测图片
            images = page.get_images(full=True)
            for img in images:
                xref = img[0]
                base_image = page.parent.extract_image(xref)
                if base_image:
                    # 获取图片位置
                    img_rect = fitz.Rect(*img[1:5])
                    block_info = {
                        "type": "image",
                        "bbox": (img_rect.x0, img_rect.y0, img_rect.x1, img_rect.y1),
                        "xref": xref,
                        "page_num": page.number + 1
                    }
                    visual_blocks.append(block_info)
            
            return visual_blocks
        except Exception as e:
            print(f"提取视觉块失败: {e}")
            return []
    
    def _is_table_candidate(self, text: str) -> bool:
        """判断是否为表格候选块"""
        # 简单判断：如果文本包含大量的制表符或对齐的数字
        lines = text.strip().split('\n')
        if len(lines) < 2:
            return False
        
        # 检查是否有类似表格的结构
        has_tab = any('\t' in line for line in lines)
        has_many_spaces = any('  ' in line for line in lines)
        
        return has_tab or has_many_spaces
    
    def _process_visual_blocks(self, blocks: list, page_num: int, file_path: Path) -> list:
        """处理视觉块"""
        processed_blocks = []
        
        for block in blocks:
            if block["type"] == "text":
                if block["content"]:
                    processed_blocks.append({
                        "type": "text",
                        "content": block["content"],
                        "bbox": block["bbox"],
                        "page_num": page_num
                    })
            
            elif block["type"] == "table_candidate":
                # 尝试使用 Camelot 提取表格
                table_content = self._extract_table(file_path, page_num)
                if table_content:
                    processed_blocks.append({
                        "type": "table",
                        "content": table_content,
                        "bbox": block["bbox"],
                        "page_num": page_num
                    })
                else:
                    # 如果表格提取失败，作为普通文本处理
                    processed_blocks.append({
                        "type": "text",
                        "content": block["content"],
                        "bbox": block["bbox"],
                        "page_num": page_num
                    })
            
            elif block["type"] == "image":
                # 生成图片描述
                image_description = self._describe_image(file_path, page_num, block["bbox"])
                processed_blocks.append({
                    "type": "image",
                    "content": image_description,
                    "bbox": block["bbox"],
                    "page_num": page_num
                })
        
        return processed_blocks
    
    def _reorder_content(self, blocks: list) -> list:
        """按阅读顺序重组内容"""
        # 按 y 坐标排序，然后按 x 坐标排序
        # 对于多栏排版，先按列分组，再按行排序
        
        if not blocks:
            return []
        
        # 检测是否为多栏排版
        columns = self._detect_columns(blocks)
        
        if len(columns) > 1:
            # 多栏排版：按列处理
            ordered_content = []
            for column in columns:
                # 按 y 坐标排序列内的块
                column_sorted = sorted(column, key=lambda b: b["bbox"][1])
                for block in column_sorted:
                    if block["content"]:
                        ordered_content.append(self._format_block_content(block))
        else:
            # 单栏排版：直接按 y 坐标排序
            sorted_blocks = sorted(blocks, key=lambda b: b["bbox"][1])
            ordered_content = [self._format_block_content(block) for block in sorted_blocks if block["content"]]
        
        return ordered_content
    
    def _detect_columns(self, blocks: list) -> list:
        """检测多栏排版"""
        if not blocks:
            return []
        
        # 提取所有文本块的 x 坐标
        x_coords = []
        for block in blocks:
            if block["type"] in ["text", "table"]:
                x_coords.append(block["bbox"][0])
        
        if not x_coords:
            return [blocks]
        
        # 使用简单的聚类算法检测列
        x_coords.sort()
        columns = []
        current_column = [x_coords[0]]
        
        for x in x_coords[1:]:
            if x - current_column[-1] < 100:  # 列间距阈值
                current_column.append(x)
            else:
                columns.append(current_column)
                current_column = [x]
        
        if current_column:
            columns.append(current_column)
        
        # 为每个列创建块列表
        column_blocks = []
        for i, column in enumerate(columns):
            column_x = sum(column) / len(column)
            column_blocks.append([])
            
            for block in blocks:
                block_x = (block["bbox"][0] + block["bbox"][2]) / 2
                # 简单的列分配
                if i == len(columns) - 1 or block_x < (column_x + columns[i+1][0]) / 2:
                    column_blocks[i].append(block)
                    break
        
        return column_blocks
    
    def _format_block_content(self, block: dict) -> str:
        """格式化块内容"""
        if block["type"] == "table":
            return f"\n[表格: 第{block['page_num']}页]\n{block['content']}\n[表格结束]"
        elif block["type"] == "image":
            return f"\n[图片: 第{block['page_num']}页]\n{block['content']}\n[图片结束]"
        else:
            return block["content"]
    
    def _extract_table(self, file_path: Path, page_num: int) -> str:
        """使用 LlamaParse 提取表格"""
        try:
            from llama_parse import LlamaParse
            import os
            
            # 检查 LlamaParse API 密钥
            api_key = os.environ.get("LLAMA_CLOUD_API_KEY")
            if not api_key:
                # 如果没有 API 密钥，使用 PyMuPDF 作为后备
                return self._extract_table_with_pymupdf(file_path, page_num)
            
            # 使用 LlamaParse 提取表格
            parser = LlamaParse(
                api_key=api_key,
                result_type="markdown",  # 直接返回 Markdown 格式
                parsing_instruction="请提取表格内容，保留完整的表格结构",
            )
            
            # 解析 PDF 文件
            documents = parser.load_data(str(file_path))
            
            if not documents:
                return ""
            
            # 提取表格内容
            table_content = []
            for i, doc in enumerate(documents):
                if "|" in doc.text:  # 检查是否包含表格
                    table_content.append(f"表格 {i+1}:\n{doc.text}")
            
            if table_content:
                return "\n".join(table_content)
            else:
                # 如果 LlamaParse 没有提取到表格，尝试使用 PyMuPDF
                return self._extract_table_with_pymupdf(file_path, page_num)
                
        except ImportError:
            # LlamaParse 未安装，使用 PyMuPDF
            return self._extract_table_with_pymupdf(file_path, page_num)
        except Exception as e:
            print(f"使用 LlamaParse 提取表格失败: {e}")
            # 失败时使用 PyMuPDF 作为后备
            return self._extract_table_with_pymupdf(file_path, page_num)
    
    def _extract_table_with_pymupdf(self, file_path: Path, page_num: int) -> str:
        """使用 PyMuPDF 提取表格（后备方法）"""
        try:
            import fitz
            
            with fitz.open(file_path) as pdf:
                page = pdf[page_num - 1]  # 页面索引从 0 开始
                
                # 尝试查找表格
                tables = page.find_tables()
                
                if not tables.tables:
                    return ""
                
                table_content = []
                for i, table in enumerate(tables.tables):
                    # 获取表格数据
                    df = table.to_pandas()
                    
                    # 检查表格是否复杂（嵌套表格）
                    if self._is_complex_table(df):
                        # 对于复杂表格，使用线性化处理
                        linearized_table = self._linearize_table(df)
                        table_content.append(f"表格 {i+1} (复杂表格):\n{linearized_table}")
                    else:
                        # 对于普通表格，转换为 Markdown 格式
                        markdown_table = self._table_to_markdown(df)
                        table_content.append(f"表格 {i+1}:\n{markdown_table}")
                
                return "\n".join(table_content)
        except Exception as e:
            print(f"使用 PyMuPDF 提取表格失败: {e}")
            return ""
    
    def _table_to_markdown(self, df) -> str:
        """将 DataFrame 转换为 Markdown 表格"""
        if df.empty:
            return ""
        
        # 构建 Markdown 表格
        markdown = []
        
        # 添加表头
        headers = df.columns.tolist()
        markdown.append("| " + " | ".join(headers) + " |")
        markdown.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        # 添加数据行
        for _, row in df.iterrows():
            row_values = [str(cell) if cell is not None else "" for cell in row]
            markdown.append("| " + " | ".join(row_values) + " |")
        
        return "\n".join(markdown)
    
    def _is_complex_table(self, df) -> bool:
        """检测表格是否复杂（嵌套表格）"""
        if df.empty:
            return False
        
        # 检查是否有合并单元格的迹象
        # 1. 检查是否有重复的行或列
        has_duplicate_rows = any(df.duplicated())
        has_duplicate_cols = any(df.columns.duplicated())
        
        # 2. 检查单元格内容是否包含表格结构
        for _, row in df.iterrows():
            for cell in row:
                if cell and isinstance(cell, str):
                    # 检查单元格内容是否包含表格特征
                    if "|" in cell and ("---" in cell or len(cell.split("|")) > 2):
                        return True
        
        # 3. 检查表格大小（过大的表格可能更复杂）
        if df.shape[0] > 20 or df.shape[1] > 10:
            return True
        
        return has_duplicate_rows or has_duplicate_cols
    
    def _linearize_table(self, df) -> str:
        """线性化处理复杂表格"""
        if df.empty:
            return ""
        
        linearized = []
        headers = df.columns.tolist()
        
        # 为每一行生成键值对描述
        for index, row in df.iterrows():
            row_desc = []
            for col, value in zip(headers, row):
                if value and str(value).strip():
                    row_desc.append(f"{col}: {value}")
            if row_desc:
                linearized.append(f"行 {index + 1}: " + ", ".join(row_desc))
        
        return "\n".join(linearized)
    
    def _describe_image(self, file_path: Path, page_num: int, bbox: tuple) -> str:
        """生成图片描述"""
        try:
            # 提取图片
            image_data = self._extract_image(file_path, page_num, bbox)
            
            if image_data:
                # 使用多模态大模型生成描述
                description = self._generate_image_description(image_data)
                if description:
                    return f"[图片描述: {description}]"
            
            # 如果无法提取或生成描述，返回占位符
            return "[图片: 第{}页包含图片]".format(page_num)
        except Exception as e:
            print(f"生成图片描述失败: {e}")
            return "[图片: 第{}页包含图片]".format(page_num)
    
    def _extract_image(self, file_path: Path, page_num: int, bbox: tuple) -> bytes:
        """提取图片数据"""
        try:
            import fitz
            
            with fitz.open(file_path) as pdf:
                page = pdf[page_num - 1]  # 页面索引从 0 开始
                
                # 获取页面上的所有图片
                images = page.get_images(full=True)
                
                for img in images:
                    xref = img[0]
                    base_image = pdf.extract_image(xref)
                    if base_image:
                        # 检查图片是否在指定的边界框内
                        img_rect = fitz.Rect(*img[1:5])
                        if img_rect.intersects(fitz.Rect(*bbox)):
                            return base_image["image"]
            
            return None
        except Exception as e:
            print(f"提取图片失败: {e}")
            return None
    
    def _generate_image_description(self, image_data: bytes) -> str:
        """使用多模态大模型生成图片描述"""
        try:
            import os
            import httpx
            
            # 检查是否配置了多模态模型 API
            # 这里可以集成 Qwen-VL 或 GPT-4o
            
            # 检查 Qwen-VL API
            qwen_api_key = os.environ.get("QWEN_API_KEY")
            if qwen_api_key:
                return self._generate_description_with_qwen(image_data, qwen_api_key)
            
            # 检查 GPT-4o API
            openai_api_key = os.environ.get("OPENAI_API_KEY")
            if openai_api_key:
                return self._generate_description_with_openai(image_data, openai_api_key)
            
            # 如果没有配置 API，返回默认描述
            return "这是一张图片，包含视觉信息"
        except Exception as e:
            print(f"生成图片描述失败: {e}")
            return "这是一张图片，包含视觉信息"
    
    def _generate_description_with_qwen(self, image_data: bytes, api_key: str) -> str:
        """使用 Qwen-VL 生成图片描述"""
        try:
            import base64
            import json
            
            # 将图片数据编码为 base64
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            
            # 构建请求
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            payload = {
                "model": "qwen-vl-plus",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "请详细描述这张图片的内容，包括图片中的主要元素、布局和可能的用途。"
                            },
                            {
                                "type": "image",
                                "image": image_base64
                            }
                        ]
                    }
                ],
                "max_tokens": 500
            }
            
            # 发送请求
            import httpx
            response = httpx.post(
                "https://dashscope.aliyuncs.com/api/v1/services/aigc/multimodal-generation/generation",
                headers=headers,
                json=payload,
                timeout=30.0
            )
            
            # 解析响应
            result = response.json()
            if "output" in result and "text" in result["output"]:
                return result["output"]["text"].strip()
            
            return "这是一张图片，包含视觉信息"
        except Exception as e:
            print(f"使用 Qwen-VL 生成描述失败: {e}")
            return "这是一张图片，包含视觉信息"
    
    def _generate_description_with_openai(self, image_data: bytes, api_key: str) -> str:
        """使用 GPT-4o 生成图片描述"""
        try:
            import base64
            import json
            
            # 将图片数据编码为 base64
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            
            # 构建请求
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            payload = {
                "model": "gpt-4o",
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "请详细描述这张图片的内容，包括图片中的主要元素、布局和可能的用途。"
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_base64}"
                                }
                            }
                        ]
                    }
                ],
                "max_tokens": 500
            }
            
            # 发送请求
            import httpx
            response = httpx.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=30.0
            )
            
            # 解析响应
            result = response.json()
            if "choices" in result and len(result["choices"]) > 0:
                return result["choices"][0]["message"]["content"].strip()
            
            return "这是一张图片，包含视觉信息"
        except Exception as e:
            print(f"使用 GPT-4o 生成描述失败: {e}")
            return "这是一张图片，包含视觉信息"
    
    def _has_images(self, page) -> bool:
        """检测页面是否包含图片"""
        try:
            if hasattr(page, 'images'):
                return len(page.images) > 0
            return False
        except:
            return False
    
    def _extract_text_with_layout(self, page) -> str:
        """提取文本并处理多栏排版"""
        try:
            # 获取页面的文本和位置信息
            words = page.extract_words()
            
            if not words:
                return ""
            
            # 检测是否为多栏排版
            if self._is_multi_column(words):
                return self._process_multi_column(words, page.width)
            else:
                return page.extract_text() or ""
        except:
            return page.extract_text() or ""
    
    def _is_multi_column(self, words) -> bool:
        """检测是否为多栏排版"""
        if not words or len(words) < 10:
            return False
        
        # 分析文本的 x 坐标分布
        x_positions = [word['x0'] for word in words]
        x_positions.sort()
        
        # 检测是否有明显的列分隔
        gaps = []
        for i in range(1, len(x_positions)):
            gap = x_positions[i] - x_positions[i-1]
            if gap > 80:  # 假设列间距大于80像素
                gaps.append(gap)
        
        # 如果有多个大间隙，可能是多栏排版
        return len(gaps) >= 1
    
    def _process_multi_column(self, words, page_width) -> str:
        """处理多栏排版"""
        try:
            # 将文本块按 x 坐标分组
            columns = {}
            
            for word in words:
                x0 = word['x0']
                x1 = word['x1']
                
                # 确定列
                column_key = int(x0 // 100)  # 每100像素为一列
                
                if column_key not in columns:
                    columns[column_key] = []
                
                columns[column_key].append(word)
            
            # 按列排序
            sorted_columns = sorted(columns.keys())
            
            # 提取每列的文本
            result = ""
            for col_key in sorted_columns:
                column_words = columns[col_key]
                # 按y坐标排序
                column_words.sort(key=lambda w: (w['top'], w['x0']))
                
                # 提取文本
                column_text = " ".join([word['text'] for word in column_words])
                result += column_text + "\n"
            
            return result
        except:
            return ""
    
    def _process_tables(self, tables) -> str:
        """处理表格"""
        try:
            table_text = ""
            
            for table_idx, table in enumerate(tables):
                if not table:
                    continue
                
                table_text += f"\n[表格 {table_idx + 1}]\n"
                
                for row_idx, row in enumerate(table):
                    if not row:
                        continue
                    
                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                    table_text += row_text + "\n"
                
                table_text += "[表格结束]\n"
            
            return table_text
        except:
            return ""
    
    def _needs_ocr(self, text: str) -> bool:
        """判断是否需要 OCR 处理"""
        if not text or not text.strip():
            return True
        
        # 检查文本长度
        if len(text) < 30:
            return True
        
        # 检查是否包含大量乱码（非 ASCII 字符）
        if len(text) > 0:
            # 对于中文文本，我们使用不同的判断标准
            # 统计中文字符数量
            chinese_chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
            
            # 如果文本主要是中文，且长度足够，则不需要 OCR
            if chinese_chars > len(text) * 0.5 and len(text) > 30:
                return False
            
            # 如果文本主要是英文，检查是否有大量非 ASCII 字符
            if chinese_chars < len(text) * 0.3:
                special_chars = sum(1 for char in text if ord(char) > 127)
                # 如果超过 70% 的字符都是非 ASCII，可能是乱码
                if special_chars / len(text) > 0.7:
                    return True
        
        # 如果文本长度足够且主要是中文或英文，则不需要 OCR
        if len(text) > 50:
            return False
        
        # 默认情况下，短文本或异常文本需要 OCR
        return True
    
    def _perform_ocr(self, file_path: Path, page_num: int) -> str:
        """执行 OCR 处理"""
        try:
            import pytesseract
            from PIL import Image
            import io
            
            # 将 PDF 页面转换为图像
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                page = reader.pages[page_num]
                
                # 尝试提取图像
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            try:
                                image_data = xObject[obj]._data
                                image = Image.open(io.BytesIO(image_data))
                                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                                return f"[OCR 处理结果]:\n{text}"
                            except:
                                continue
                
                # 如果无法直接提取图像，使用 pdf2image
                try:
                    from pdf2image import convert_from_path
                    images = convert_from_path(str(file_path), first_page=page_num+1, last_page=page_num+1)
                    if images:
                        text = pytesseract.image_to_string(images[0], lang='chi_sim+eng')
                        return f"[OCR 处理结果]:\n{text}"
                except:
                    pass
                
                return "[OCR 处理失败: 无法提取图像]"
                
        except ImportError:
            return "[OCR 处理失败: 需要安装 pytesseract 和 pdf2image 库]"
        except Exception as e:
            return f"[OCR 处理失败: {str(e)}]"
    
    def _parse_word(self, file_path: Path) -> str:
        """解析 Word 文件"""
        try:
            import docx
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise Exception("解析 Word 文件需要 python-docx 库，请安装: pip install python-docx")
        except Exception as e:
            raise Exception(f"解析 Word 文件时出错: {str(e)}")
    
    def _parse_json(self, file_path: Path) -> str:
        """解析 JSON 文件"""
        try:
            import json
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            # 将 JSON 转换为文本格式
            return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            raise Exception(f"解析 JSON 文件时出错: {str(e)}")
